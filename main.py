import os
import streamlit as st
import google.generativeai as genai
import pandas as pd
from docx import Document
import io
import re
from typing import Tuple, List, Optional
import logging
from pathlib import Path
from llama_parse import LlamaParse
from dotenv import load_dotenv
import tempfile

# Load environment variables
load_dotenv()

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Get API keys from environment variables
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")
LLAMA_CLOUD_API_KEY = os.getenv("LLAMA_CLOUD_API_KEY")

if not GEMINI_API_KEY or not LLAMA_CLOUD_API_KEY:
    raise ValueError("Missing required API keys in environment variables")

# Configure Gemini API
genai.configure(api_key=GEMINI_API_KEY)

# Initialize LlamaParse
parser = LlamaParse(
    api_key=LLAMA_CLOUD_API_KEY,
    result_type="markdown",
    verbose=True
)

# Create the model
generation_config = {
    "temperature": 0.7,
    "top_p": 0.95,
    "top_k": 64,
    "max_output_tokens": 8192,
    "response_mime_type": "text/plain",
}

safety_settings = [
    {"category": "HARM_CATEGORY_HARASSMENT", "threshold": "BLOCK_MEDIUM_AND_ABOVE"},
    {"category": "HARM_CATEGORY_HATE_SPEECH", "threshold": "BLOCK_MEDIUM_AND_ABOVE"},
    {"category": "HARM_CATEGORY_SEXUALLY_EXPLICIT", "threshold": "BLOCK_MEDIUM_AND_ABOVE"},
    {"category": "HARM_CATEGORY_DANGEROUS_CONTENT", "threshold": "BLOCK_MEDIUM_AND_ABOVE"},
]

# Update model initialization
try:
    model = genai.GenerativeModel(
        model_name="gemini-2.0-flash-thinking-exp-1219",  # Using stable model version
        generation_config=generation_config,
        safety_settings=safety_settings,
    )
except Exception as e:
    logger.error(f"Failed to initialize Gemini model: {e}")
    raise

def extract_text_from_pdf(file) -> Tuple[str, List[pd.DataFrame]]:
    """
    Extract text from PDF file using LlamaParse
    
    Args:
        file: File-like object containing PDF
        
    Returns:
        Tuple[str, List[pd.DataFrame]]: Extracted text and tables
    """
    try:
        # Create a temporary file to save the uploaded content
        with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as tmp_file:
            tmp_file.write(file.getvalue())
            tmp_path = tmp_file.name

        # Parse the PDF using LlamaParse
        result = parser.load_data(tmp_path)
        
        # Combine all pages into one text
        text = "\n\n---\n\n".join([page.text for page in result])
        
        # Extract tables if available
        tables = []
        for page in result:
            if hasattr(page, 'tables') and page.tables:
                for table in page.tables:
                    try:
                        df = pd.DataFrame(table)
                        if not df.empty:
                            tables.append(df)
                    except Exception as e:
                        logger.warning(f"Failed to convert table to DataFrame: {e}")

        # Clean up temporary file
        os.unlink(tmp_path)
            
        return text, tables
    except Exception as e:
        logger.error(f"PDF extraction failed: {e}")
        raise RuntimeError(f"Failed to process PDF: {str(e)}")

def extract_text_from_docx(file):
    """Extract text and tables from DOCX file"""
    doc = Document(file)
    text = ""
    tables = []
    
    # Extract text
    for paragraph in doc.paragraphs:
        text += paragraph.text + "\n"
    
    # Extract tables
    for table in doc.tables:
        table_data = []
        for row in table.rows:
            row_data = [cell.text for cell in row.cells]
            table_data.append(row_data)
        tables.append(pd.DataFrame(table_data[1:], columns=table_data[0]))
    
    return text, tables

def extract_text_from_txt(file):
    """Extract text from TXT file"""
    text = file.getvalue().decode()
    return text, []

def clean_text(text):
    """Clean extracted text"""
    # Remove excessive whitespace
    text = re.sub(r'\s+', ' ', text)
    # Remove special characters but keep basic punctuation
    text = re.sub(r'[^\w\s.,!?-]', '', text)
    return text.strip()

def analyze_section(text, section_type="requirements"):
    """Analyze specific sections of the RFP"""
    chat = model.start_chat(history=[])
    
    prompts = {
        "requirements": """
            Analyze the technical requirements in this section:
            1. List all mandatory requirements
            2. Identify optional requirements
            3. Flag any unclear specifications
            4. Note any technical constraints
            """,
        "timeline": """
            Analyze the timeline aspects:
            1. List all key dates and deadlines
            2. Identify major milestones
            3. Flag any tight or unrealistic timelines
            4. Note dependencies between phases
            """,
        "compliance": """
            Analyze compliance requirements:
            1. List all regulatory requirements
            2. Identify certification needs
            3. Note security requirements
            4. Flag critical compliance issues
            """
    }
    
    response = chat.send_message(f"{prompts.get(section_type, prompts['requirements'])}\n\nContent:\n{text}")
    return response.text

def analyze_tables(tables):
    """Analyze tables found in the document"""
    analysis = []
    for i, table in enumerate(tables):
        if not table.empty:
            analysis.append(f"\nTable {i+1} Analysis:")
            analysis.append(f"- Columns: {', '.join(table.columns.tolist())}")
            analysis.append(f"- Rows: {len(table)}")
            analysis.append("- Content Summary:")
            for col in table.columns:
                unique_values = table[col].nunique()
                analysis.append(f"  * {col}: {unique_values} unique values")
    
    return "\n".join(analysis)

def analyze_rfp_content(text: str) -> str:
    """
    Analyze the full content of the RFP document.
    """
    chat = model.start_chat(history=[])
    prompt = """
    Analyze this RFP document and provide a detailed markdown-formatted response covering:

    # Executive Summary
    - Brief overview of the RFP
    - Key objectives and goals

    # Requirements Analysis
    - ## Technical Requirements
      - Core technical specifications
      - Optional features
      - Integration requirements
    - ## Business Requirements
      - Mandatory business needs
      - Optional enhancements
    
    # Timeline and Milestones
    - Key dates and deadlines
    - Project phases
    - Dependencies

    # Risk Assessment
    - Technical risks
    - Business risks
    - Compliance concerns

    # Recommendations
    - Strategic approach
    - Key focus areas
    - Potential challenges

    Please analyze the following content:
    {text}
    """
    try:
        response = chat.send_message(prompt.format(text=text))
        return response.text
    except Exception as e:
        logger.error(f"Error analyzing RFP content: {e}")
        return "Error analyzing content. Please try again."

# Update the main function to include caching
@st.cache_data
def process_document(file_content, file_type: str) -> Tuple[str, List[pd.DataFrame]]:
    """Cache the document processing results"""
    if file_type == 'pdf':
        return extract_text_from_pdf(file_content)
    elif file_type == 'docx':
        return extract_text_from_docx(file_content)
    else:  # txt
        return extract_text_from_txt(file_content)

def main():
    st.set_page_config(page_title="RFP Document Analyzer", layout="wide")
    
    # Check for API keys
    if not GEMINI_API_KEY or not LLAMA_CLOUD_API_KEY:
        st.error("Missing required API keys in environment variables!")
        st.stop()
    
    # Move controls to sidebar
    with st.sidebar:
        st.title("📄 RFP Document Analyzer")
        st.markdown("Upload your RFP document for analysis.")
        
        # File uploader with size limit
        MAX_FILE_SIZE = 50 * 1024 * 1024  # 50MB
        uploaded_file = st.file_uploader(
            "Choose an RFP document",
            type=['pdf', 'docx', 'txt'],
            help="Upload a PDF, DOCX, or TXT file (max 50MB)"
        )
        
        if uploaded_file:
            st.write(f"File: {uploaded_file.name}")
            st.write(f"Size: {uploaded_file.size/1024:.2f} KB")
            
            # Analysis control buttons
            analyze_full = st.button("🔍 Analyze Full Content", use_container_width=True)
            analyze_req = st.button("📋 Analyze Requirements", use_container_width=True)
            analyze_timeline = st.button("⏱️ Analyze Timeline", use_container_width=True)
            analyze_compliance = st.button("✓ Analyze Compliance", use_container_width=True)
        
        # Add tips in sidebar
        st.markdown("---")
        st.markdown("### 💡 Analysis Tips")
        st.markdown("""
        1. Upload your RFP document
        2. Choose analysis type
        3. Review results in main panel
        """)

    # Main content area
    if uploaded_file:
        if uploaded_file.size > MAX_FILE_SIZE:
            st.error("File size exceeds 50MB limit. Please upload a smaller file.")
            return
            
        try:
            # Process document with caching
            text, tables = process_document(
                uploaded_file,
                uploaded_file.name.split('.')[-1].lower()
            )
            
            # Clean extracted text
            cleaned_text = clean_text(text)
            
            # Create two columns for better layout
            col1, col2 = st.columns([2, 1])
            
            with col1:
                st.markdown("### 📑 Analysis Results")
                
                # Show analysis based on button clicks
                if analyze_full:
                    with st.spinner("Analyzing full content..."):
                        analysis = analyze_rfp_content(cleaned_text)
                        st.markdown(analysis)
                
                if analyze_req:
                    with st.spinner("Analyzing requirements..."):
                        req_analysis = analyze_section(cleaned_text, "requirements")
                        st.markdown(req_analysis)
                
                if analyze_timeline:
                    with st.spinner("Analyzing timeline..."):
                        timeline_analysis = analyze_section(cleaned_text, "timeline")
                        st.markdown(timeline_analysis)
                
                if analyze_compliance:
                    with st.spinner("Analyzing compliance..."):
                        compliance_analysis = analyze_section(cleaned_text, "compliance")
                        st.markdown(compliance_analysis)
            
            with col2:
                # Show extracted text and tables
                with st.expander("📄 Extracted Text", expanded=False):
                    st.text_area("Content", cleaned_text, height=300)
                
                if tables:
                    with st.expander("📊 Extracted Tables", expanded=False):
                        table_analysis = analyze_tables(tables)
                        st.markdown(table_analysis)
                        for i, table in enumerate(tables):
                            st.markdown(f"**Table {i+1}**")
                            st.dataframe(table, use_container_width=True)
                
        except Exception as e:
            st.error(f"Error processing document: {str(e)}")
            st.error("Please make sure the document is not corrupted and try again.")
    else:
        # Show welcome message when no file is uploaded
        st.markdown("""
        # Welcome to RFP Document Analyzer! 👋
        
        This tool helps you analyze Request for Proposal (RFP) documents by:
        
        - Extracting and analyzing content
        - Identifying key requirements
        - Analyzing timelines and milestones
        - Checking compliance requirements
        
        To get started, upload your RFP document using the sidebar.
        """)

if __name__ == "__main__":
    main()
